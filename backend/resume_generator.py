"""Generate tailored resumes as PDF and DOCX from structured data."""

import os
import json
import zipfile
import logging
from datetime import datetime
from io import BytesIO

logger = logging.getLogger(__name__)


RESUME_TEMPLATE_CSS = """
@page {
    size: letter;
    margin: 0.6in 0.7in;
}
body {
    font-family: Arial, Helvetica, sans-serif;
    font-size: 10.5pt;
    line-height: 1.35;
    color: #1a1a1a;
    margin: 0;
    padding: 0;
}
.header {
    text-align: center;
    border-bottom: 2px solid #2c3e50;
    padding-bottom: 8px;
    margin-bottom: 12px;
}
.header h1 {
    font-size: 18pt;
    margin: 0 0 2px 0;
    color: #2c3e50;
    letter-spacing: 0.5px;
}
.header .contact {
    font-size: 9.5pt;
    color: #555;
}
.section {
    margin-bottom: 10px;
}
.section h2 {
    font-size: 11.5pt;
    color: #2c3e50;
    text-transform: uppercase;
    letter-spacing: 1px;
    border-bottom: 1px solid #bdc3c7;
    padding-bottom: 3px;
    margin: 0 0 6px 0;
}
.summary {
    font-size: 10pt;
    color: #333;
    line-height: 1.4;
}
.job-entry {
    margin-bottom: 8px;
}
.job-header {
    display: flex;
    justify-content: space-between;
    align-items: baseline;
}
.job-title {
    font-weight: bold;
    font-size: 10.5pt;
}
.job-date {
    font-size: 9.5pt;
    color: #666;
    white-space: nowrap;
}
.job-company {
    font-style: italic;
    color: #444;
    font-size: 10pt;
}
.job-location {
    font-size: 9.5pt;
    color: #666;
}
ul {
    margin: 3px 0 0 0;
    padding-left: 18px;
}
li {
    margin-bottom: 2px;
}
.skills-list {
    margin-top: 4px;
}
.skill-category {
    margin-bottom: 6px;
    padding-left: 15px;
}
.skill-name {
    font-weight: bold;
    color: #2c3e50;
    font-size: 10.5pt;
    margin-bottom: 1px;
}
.skill-details {
    padding-left: 15px;
    font-size: 9.5pt;
    color: #333;
    line-height: 1.3;
}
.edu-entry {
    margin-bottom: 8px;
}
.edu-header {
    display: flex;
    justify-content: space-between;
    align-items: baseline;
}
.edu-title {
    font-weight: bold;
    font-size: 10.5pt;
}
.edu-date {
    font-size: 9.5pt;
    color: #666;
    white-space: nowrap;
}
.edu-institution {
    font-style: italic;
    color: #444;
    font-size: 10pt;
}
"""


def _build_html(name, contact, summary, experience_list, education_list, skills_categories):
    """Build HTML resume from structured data.

    Args:
        name: Full name string.
        contact: Contact info string (email, phone, location, linkedin).
        summary: Professional summary text.
        experience_list: List of dicts with keys: title, company, location,
                         start_date, end_date, bullets (list of strings).
        education_list: List of dicts with keys: institution, degree,
                        field_of_study, gpa, dates.
        skills_categories: List of dicts with keys: category_name, skills.
    """
    contact_info, contact_links = contact if isinstance(contact, tuple) else (contact, '')
    contact_html = f"<div>{contact_info}</div>"
    if contact_links:
        contact_html += f"<div>{contact_links}</div>"

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>{RESUME_TEMPLATE_CSS}</style>
</head>
<body>
<div class="header">
    <h1>{name}</h1>
    <div class="contact">{contact_html}</div>
</div>
"""

    if summary:
        html += f"""<div class="section">
    <h2>Professional Summary</h2>
    <div class="summary">{summary}</div>
</div>
"""

    # Experience
    if experience_list:
        html += '<div class="section"><h2>Experience</h2>\n'
        for exp in experience_list:
            title = exp.get('title', '')
            company = exp.get('company', '')
            location = exp.get('location', '')
            start_date = exp.get('start_date', '')
            end_date = exp.get('end_date', '')

            # Build date range string
            date_parts = [p for p in (start_date, end_date) if p]
            date_str = ' - '.join(date_parts) if date_parts else ''

            # Build company-location string
            company_loc = company
            if location:
                company_loc = f"{company}, {location}"

            html += '<div class="job-entry">\n'
            html += '  <div class="job-header">\n'
            html += f'    <div class="job-title">{title}</div>\n'
            if date_str:
                html += f'    <div class="job-date">{date_str}</div>\n'
            html += '  </div>\n'
            html += f'  <div class="job-company">{company_loc}</div>\n'

            bullets = exp.get('bullets', [])
            if bullets:
                html += '<ul>\n'
                for bullet in bullets:
                    html += f'  <li>{bullet}</li>\n'
                html += '</ul>\n'
            html += '</div>\n'
        html += '</div>\n'

    # Skills
    if skills_categories:
        html += '<div class="section"><h2>Skills</h2>\n'
        for cat in skills_categories:
            cat_name = cat.get('category_name', '')
            skills_text = cat.get('skills', '')
            if cat_name:
                html += f'<div class="skill-category">\n      <div class="skill-name">{cat_name}</div>\n      <div class="skill-details">{skills_text}</div>\n    </div>\n'
            else:
                html += f'<div class="skill-category">\n      <div class="skill-details">{skills_text}</div>\n    </div>\n'
        html += '</div>\n'

    # Education
    if education_list:
        html += '<div class="section"><h2>Education</h2>\n'
        for edu in education_list:
            institution = edu.get('institution', '')
            degree = edu.get('degree', '')
            field = edu.get('field_of_study', '')
            gpa = edu.get('gpa', '')
            dates = edu.get('dates', '')

            parts = []
            if degree and field:
                parts.append(f"{degree} in {field}")
            elif degree:
                parts.append(degree)
            elif field:
                parts.append(field)
            degree_str = ', '.join(parts)

            inst_str = institution
            if gpa:
                inst_str += f" (GPA: {gpa})"

            html += '  <div class="edu-entry">\n'
            html += '    <div class="edu-header">\n'
            html += f'      <div class="edu-title">{degree_str}</div>\n'
            if dates:
                html += f'      <div class="edu-date">{dates}</div>\n'
            html += '    </div>\n'
            html += f'    <div class="edu-institution">{inst_str}</div>\n'
            html += '  </div>\n'
        html += '</div>\n'

    html += '</body></html>'
    return html


def generate_pdf(html_content):
    """Generate PDF from HTML using weasyprint."""
    try:
        from weasyprint import HTML
        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes
    except ImportError:
        return None


def generate_docx(name, contact, summary, experience_list, education_list, skills_categories):
    """Generate DOCX using python-docx from structured data."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    except ImportError:
        return None

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10.5)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

    # Contact
    if contact:
        contact_info, contact_links = contact if isinstance(contact, tuple) else (contact, '')
        if contact_info:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2) if contact_links else Pt(8)
            run = p.add_run(contact_info)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        if contact_links:
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(8)
            run2 = p2.add_run(contact_links)
            run2.font.size = Pt(9.5)
            run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Summary
    if summary:
        h = doc.add_heading('Professional Summary', level=2)
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        for run in h.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15

    # Experience
    if experience_list:
        h = doc.add_heading('Experience', level=2)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        for run in h.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

        for exp in experience_list:
            title = exp.get('title', '')
            company = exp.get('company', '')
            location = exp.get('location', '')
            start_date = exp.get('start_date', '')
            end_date = exp.get('end_date', '')

            date_parts = [p for p in (start_date, end_date) if p]
            date_str = ' - '.join(date_parts)

            company_loc = company
            if location:
                company_loc = f"{company}, {location}"

            # Title & Date line (using right-aligned tab stop)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
            
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(10.5)
            
            if date_str:
                run_date = p.add_run(f"\t{date_str}")
                run_date.font.size = Pt(9.5)
                run_date.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Company line
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(4)
            p2.paragraph_format.keep_with_next = True
            run = p2.add_run(company_loc)
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            # Bullets
            for bullet in exp.get('bullets', []):
                bp = doc.add_paragraph(bullet, style='List Bullet')
                bp.paragraph_format.space_before = Pt(0)
                bp.paragraph_format.space_after = Pt(2)
                for run in bp.runs:
                    run.font.name = 'Arial'

    # Skills
    if skills_categories:
        h = doc.add_heading('Skills', level=2)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        for run in h.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

        for cat in skills_categories:
            cat_name = cat.get('category_name', '')
            skills_text = cat.get('skills', '')
            if cat_name:
                p1 = doc.add_paragraph()
                p1.paragraph_format.space_before = Pt(2)
                p1.paragraph_format.space_after = Pt(1)
                p1.paragraph_format.left_indent = Inches(0.2)
                p1.paragraph_format.keep_with_next = True
                run = p1.add_run(cat_name)
                run.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(4)
                p2.paragraph_format.left_indent = Inches(0.4)
                run = p2.add_run(skills_text)
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            else:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(4)
                p2.paragraph_format.left_indent = Inches(0.2)
                run = p2.add_run(skills_text)
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)

    # Education
    if education_list:
        h = doc.add_heading('Education', level=2)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        for run in h.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
            
        for edu in education_list:
            institution = edu.get('institution', '')
            degree = edu.get('degree', '')
            field = edu.get('field_of_study', '')
            gpa = edu.get('gpa', '')
            dates = edu.get('dates', '')

            parts = []
            if degree and field:
                parts.append(f"{degree} in {field}")
            elif degree:
                parts.append(degree)
            elif field:
                parts.append(field)
            degree_str = ', '.join(parts)

            inst_str = institution
            if gpa:
                inst_str += f"  |  GPA: {gpa}"

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
            run = p.add_run(degree_str)
            run.bold = True
            run.font.size = Pt(10.5)
            if dates:
                run_date = p.add_run(f"\t{dates}")
                run_date.font.size = Pt(9.5)
                run_date.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(6)
            run = p2.add_run(inst_str)
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def create_zip(pdf_bytes, docx_bytes, job_id):
    """Create a zip file containing both PDF and DOCX."""
    buf = BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        if pdf_bytes:
            zf.writestr(f'michael_sounhein_resume_job_{job_id}.pdf', pdf_bytes)
        if docx_bytes:
            zf.writestr(f'michael_sounhein_resume_job_{job_id}.docx', docx_bytes)
    buf.seek(0)
    return buf.read()


def generate_resume(job_id, output_dir, summary_override='', experience_order=None, skills_emphasis=None, bullet_overrides=None):
    """
    Generate tailored resume files from structured database tables.
    NEVER mutates the master tables — works on copies only.

    Args:
        job_id: Job ID for file naming.
        output_dir: Directory to write output files.
        summary_override: Override the profile summary if provided.
        experience_order: Optional list of experience IDs to reorder entries.
        skills_emphasis: Optional list of skill category names to prioritize.
        bullet_overrides: Optional dict mapping experience ID -> list of bullet
            indices to INCLUDE. If None for an experience, all bullets are kept.
            Example: {1: [0, 2, 5], 3: [0, 1]} means "for exp 1 keep bullets
            0,2,5; for exp 3 keep bullets 0,1; for all others keep everything."

    Returns:
        dict with keys: html, pdf_path, docx_path, zip_path
    """
    from database import SessionLocal
    from models import ResumeProfile, ResumeExperience, ResumeEducation, ResumeSkillCategory

    db = SessionLocal()
    try:
        # Load profile
        profile = db.query(ResumeProfile).first()
        name = profile.full_name if profile and profile.full_name else 'Your Name'

        # Build contact string
        contact_info_parts = []
        contact_links_parts = []
        if profile:
            if profile.email:
                contact_info_parts.append(profile.email)
            if profile.phone:
                contact_info_parts.append(profile.phone)
            if profile.location:
                contact_info_parts.append(profile.location)
            if profile.linkedin_url:
                contact_links_parts.append(profile.linkedin_url)
            if profile.github_url:
                contact_links_parts.append(profile.github_url)
        contact = (' | '.join(contact_info_parts), ' | '.join(contact_links_parts))

        # Summary
        summary = summary_override
        if not summary and profile and profile.summary:
            summary = profile.summary

        # Experience
        exp_rows = db.query(ResumeExperience).order_by(ResumeExperience.sort_order).all()
        experience_list = []
        for e in exp_rows:
            experience_list.append({
                'id': e.id,
                'title': e.title,
                'company': e.company,
                'location': e.location,
                'start_date': e.start_date,
                'end_date': e.end_date,
                'bullets': json.loads(e.bullets) if e.bullets else [],
            })

        # Apply experience_order if provided (list of experience IDs in desired order)
        if experience_order:
            id_to_exp = {exp['id']: exp for exp in experience_list}
            reordered = []
            for eid in experience_order:
                if eid in id_to_exp:
                    reordered.append(id_to_exp.pop(eid))
            # Append any remaining
            reordered.extend(id_to_exp.values())
            experience_list = reordered

        # Apply bullet_overrides — filter which bullets to keep per experience
        # This operates on the in-memory copy, NEVER touches the DB
        if bullet_overrides:
            for exp in experience_list:
                exp_id = exp['id']
                if exp_id in bullet_overrides:
                    keep_indices = bullet_overrides[exp_id]
                    exp['bullets'] = [exp['bullets'][i] for i in keep_indices if i < len(exp['bullets'])]

        # Education
        edu_rows = db.query(ResumeEducation).order_by(ResumeEducation.sort_order).all()
        education_list = []
        for e in edu_rows:
            education_list.append({
                'institution': e.institution,
                'degree': e.degree,
                'field_of_study': e.field_of_study,
                'gpa': e.gpa,
                'dates': e.dates,
            })

        # Skills
        skill_rows = db.query(ResumeSkillCategory).order_by(ResumeSkillCategory.sort_order).all()
        skills_categories = []
        for s in skill_rows:
            skills_categories.append({
                'category_name': s.category_name,
                'skills': s.skills,
            })

        # Apply skills_emphasis if provided (move matching categories to front)
        if skills_emphasis:
            emphasized = []
            rest = []
            for cat in skills_categories:
                if cat['category_name'].lower() in [e.lower() for e in skills_emphasis]:
                    emphasized.append(cat)
                else:
                    rest.append(cat)
            skills_categories = emphasized + rest

    finally:
        db.close()

    # Build HTML
    html = _build_html(name, contact, summary, experience_list, education_list, skills_categories)

    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    # Generate PDF
    pdf_path = os.path.join(output_dir, f'resume_{job_id}_{timestamp}.pdf')
    pdf_bytes = generate_pdf(html)
    if pdf_bytes:
        with open(pdf_path, 'wb') as f:
            f.write(pdf_bytes)
    else:
        pdf_path = None

    # Generate DOCX
    docx_path = os.path.join(output_dir, f'resume_{job_id}_{timestamp}.docx')
    docx_bytes = generate_docx(name, contact, summary, experience_list, education_list, skills_categories)
    if docx_bytes:
        with open(docx_path, 'wb') as f:
            f.write(docx_bytes)
    else:
        docx_path = None

    # Generate ZIP
    zip_path = os.path.join(output_dir, f'resume_{job_id}_{timestamp}.zip')
    zip_bytes = create_zip(pdf_bytes, docx_bytes, job_id)
    with open(zip_path, 'wb') as f:
        f.write(zip_bytes)

    return {
        'html': html,
        'pdf_path': pdf_path,
        'docx_path': docx_path,
        'zip_path': zip_path,
    }
